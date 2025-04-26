# document_handler.py
import os
from PyPDF2 import PdfReader, PdfWriter
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import logging
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import tempfile

# Configure logger
logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def read_pdf(file_path):
    """Чтение текста из PDF файла."""
    try:
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        logger.error(f"Ошибка чтения PDF: {e}")
        return None

def read_docx(file_path):
    """Чтение текста из DOCX файла."""
    try:
        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Ошибка чтения DOCX: {e}")
        return None

def read_txt(file_path):
    """Чтение текста из TXT файла."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Ошибка чтения TXT: {e}")
        return None

def read_py(file_path):
    """Чтение текста из PY файла."""
    return read_txt(file_path)

def generate_document(content, file_path, file_type='txt', title=None, author=None, formatting=None):
    """
    Генерация документа заданного типа с поддержкой форматирования.
    
    Args:
        content (str): Содержимое документа
        file_path (str): Путь для сохранения файла
        file_type (str): Тип файла (txt, docx, pdf)
        title (str, optional): Заголовок документа
        author (str, optional): Автор документа
        formatting (dict, optional): Настройки форматирования:
            - font_size (int): Размер шрифта
            - font_name (str): Название шрифта
            - line_spacing (float): Межстрочный интервал
            - margins (dict): Отступы (left, right, top, bottom)
            - text_color (tuple): Цвет текста (R, G, B)
            - alignment (str): Выравнивание (left, center, right, justify)
    """
    try:
        if file_type == 'txt':
            with open(file_path, 'w', encoding='utf-8') as file:
                if title:
                    file.write(f"{title}\n{'=' * len(title)}\n\n")
                if author:
                    file.write(f"Автор: {author}\n\n")
                file.write(content)
                
        elif file_type == 'docx':
            doc = DocxDocument()
            
            # Настройки форматирования по умолчанию
            formatting = formatting or {}
            font_size = formatting.get('font_size', 12)
            font_name = formatting.get('font_name', 'Times New Roman')
            line_spacing = formatting.get('line_spacing', 1.5)
            margins = formatting.get('margins', {'left': 1, 'right': 1, 'top': 1, 'bottom': 1})
            text_color = formatting.get('text_color', (0, 0, 0))
            alignment = formatting.get('alignment', 'left')
            
            # Установка отступов
            sections = doc.sections
            for section in sections:
                section.left_margin = Pt(margins['left'] * 72)
                section.right_margin = Pt(margins['right'] * 72)
                section.top_margin = Pt(margins['top'] * 72)
                section.bottom_margin = Pt(margins['bottom'] * 72)
            
            # Добавление заголовка
            if title:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title)
                title_run.font.size = Pt(font_size + 4)
                title_run.font.bold = True
                title_run.font.color.rgb = RGBColor(*text_color)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавление автора
            if author:
                author_para = doc.add_paragraph()
                author_run = author_para.add_run(f"Автор: {author}")
                author_run.font.size = Pt(font_size)
                author_run.font.italic = True
                author_run.font.color.rgb = RGBColor(*text_color)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавление основного текста
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(content)
            content_run.font.size = Pt(font_size)
            content_run.font.name = font_name
            content_run.font.color.rgb = RGBColor(*text_color)
            
            # Установка межстрочного интервала
            content_para.paragraph_format.line_spacing = line_spacing
            
            # Установка выравнивания
            if alignment == 'left':
                content_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif alignment == 'center':
                content_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                content_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'justify':
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            doc.save(file_path)
            
        elif file_type == 'pdf':
            # Создаем временный файл для PDF
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_path = temp_file.name
            temp_file.close()
            
            # Создаем PDF документ
            c = canvas.Canvas(temp_path, pagesize=letter)
            
            # Настройки форматирования по умолчанию
            formatting = formatting or {}
            font_size = formatting.get('font_size', 12)
            font_name = formatting.get('font_name', 'Helvetica')
            line_spacing = formatting.get('line_spacing', 1.5)
            margins = formatting.get('margins', {'left': 1, 'right': 1, 'top': 1, 'bottom': 1})
            text_color = formatting.get('text_color', (0, 0, 0))
            
            # Установка шрифта
            c.setFont(font_name, font_size)
            
            # Установка цвета
            c.setFillColorRGB(*text_color)
            
            # Расчет отступов в точках
            left_margin = margins['left'] * 72
            right_margin = margins['right'] * 72
            top_margin = margins['top'] * 72
            bottom_margin = margins['bottom'] * 72
            
            # Расчет ширины и высоты текстовой области
            width = letter[0] - left_margin - right_margin
            height = letter[1] - top_margin - bottom_margin
            
            # Добавление заголовка
            if title:
                c.setFont(font_name, font_size + 4)
                c.drawCentredString(letter[0]/2, letter[1] - top_margin - font_size - 4, title)
                c.setFont(font_name, font_size)
            
            # Добавление автора
            if author:
                c.drawCentredString(letter[0]/2, letter[1] - top_margin - 2*(font_size + 4), f"Автор: {author}")
            
            # Добавление основного текста
            text_object = c.beginText(left_margin, letter[1] - top_margin - 3*(font_size + 4))
            text_object.setFont(font_name, font_size)
            text_object.setLeading(font_size * line_spacing)
            
            # Разбиваем текст на строки
            lines = content.split('\n')
            for line in lines:
                text_object.textLine(line)
            
            c.drawText(text_object)
            c.save()
            
            # Перемещаем временный файл в нужное место
            os.replace(temp_path, file_path)
            
        else:
            raise ValueError(f"Неподдерживаемый тип файла: {file_type}")
            
        return True
        
    except Exception as e:
        logger.error(f"Ошибка генерации документа: {e}")
        return False
